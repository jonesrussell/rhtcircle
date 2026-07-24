#!/usr/bin/env python3
"""Build the public, editable Sagamok member concern letter template."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "F4F6F9"


def set_run_font(run, size: float = 11, color: RGBColor = BLACK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top: int = 120, start: int = 160, bottom: int = 120, end: int = 160) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(paragraph, after=0, line=1.0)
    left = paragraph.add_run("Sagamok member letter template")
    set_run_font(left, size=9, color=MUTED, bold=True)
    right = paragraph.add_run("\tPrepared by members, for members")
    set_run_font(right, size=9, color=MUTED)

    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.5), alignment=2)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(paragraph, before=0, after=0, line=1.0)
    run = paragraph.add_run("RHT Circle | rhtcircle.ca")
    set_run_font(run, size=8.5, color=MUTED)


def add_metadata_line(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=2, line=1.0)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, bold=True)
    run = paragraph.add_run(value)
    set_run_font(run)


def add_numbered_prompt(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)


def build(output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    set_paragraph(title, before=0, after=4, line=1.0)
    run = title.add_run("MEMBER CONCERN AND RECORDS REQUEST")
    set_run_font(run, size=23, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, before=0, after=14, line=1.1)
    run = subtitle.add_run("Editable letter template for Sagamok members")
    set_run_font(run, size=13, color=MUTED)

    add_metadata_line(doc, "Date", "[Date]")
    add_metadata_line(doc, "To", "Sagamok Chief and Council")
    add_metadata_line(doc, "Cc", "Political Office and Council Secretariat")
    add_metadata_line(doc, "Subject", "[Short description of your concern]")

    spacer = doc.add_paragraph()
    set_paragraph(spacer, before=0, after=4, line=1.0)

    paragraph = doc.add_paragraph("Aanii,")
    set_paragraph(paragraph)

    paragraph = doc.add_paragraph()
    set_paragraph(paragraph)
    paragraph.add_run("I am a Sagamok member writing about ")
    prompt = paragraph.add_run("[name the service, decision, policy, meeting, or record].")
    set_run_font(prompt, color=DARK_BLUE, italic=True)

    heading = doc.add_paragraph("What happened, or what I am concerned about", style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=10)
    prompt = paragraph.add_run("[Describe what you personally experienced, what decision you heard about, or what information is missing. Add dates when you know them.]")
    set_run_font(prompt, color=DARK_BLUE, italic=True)

    doc.add_paragraph("I am asking for", style="Heading 2")
    add_numbered_prompt(doc, "[The written answer, record, decision, or action you want.]")
    add_numbered_prompt(doc, "[Any second question or document you need.]")
    add_numbered_prompt(doc, "[Who is responsible for the follow-up and when it will be completed.]")

    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=4, after=10)
    paragraph.add_run("Please confirm that this letter was received. Please reply in writing by ")
    prompt = paragraph.add_run("[date].")
    set_run_font(prompt, color=DARK_BLUE, italic=True)

    paragraph = doc.add_paragraph("Miigwech,")
    set_paragraph(paragraph, after=3)
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=2)
    prompt = paragraph.add_run("[Your name]")
    set_run_font(prompt, color=DARK_BLUE, italic=True)
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=12)
    prompt = paragraph.add_run("[Your email or phone, if you want to include it]")
    set_run_font(prompt, color=DARK_BLUE, italic=True)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, after=4)
    run = paragraph.add_run("Before sending")
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)
    for text in (
        "Replace every bracketed prompt with your own words, or delete it.",
        "Keep your sent email or a copy of the letter.",
        "If you do not receive an answer, follow up using the date of your first letter.",
    ):
        paragraph = cell.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.38)
        paragraph.paragraph_format.first_line_indent = Inches(-0.19)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.167
        run = paragraph.add_run(text)
        set_run_font(run, size=10)

    doc.core_properties.title = "Sagamok Member Concern and Records Request Letter Template"
    doc.core_properties.subject = "Editable generic member letter template"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "Sagamok, member letter, records request"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)


if __name__ == "__main__":
    main()
